from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = docx.oxml.OxmlElement(tag)
                tcPr.append(element)
            
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return heading

def add_paragraph_custom(doc, text, bold=False, size=11):
    """添加自定义段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.font.bold = bold
    return p

def create_doc():
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)
    
    # 标题
    title = doc.add_heading('人选筛选助手 Chrome Extension', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.color.rgb = RGBColor(0, 102, 204)
    
    # 副标题
    subtitle = doc.add_paragraph('使用说明书')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    
    # 一、插件简介
    add_heading_custom(doc, '一、插件简介', 1)
    
    p = add_paragraph_custom(doc, '人选筛选助手是一款专为招聘场景设计的 Chrome 浏览器插件，能够自动从 Moka 招聘系统中提取候选人信息，并通过大语言模型进行智能筛选，大幅提高简历筛选效率。')
    
    add_heading_custom(doc, '主要功能', 2)
    
    features = [
        ('🤖 AI 智能筛选', '使用大语言模型自动分析候选人简历'),
        ('📄 简历自动抓取', '自动提取候选人详情页简历图片'),
        ('📊 批量处理', '支持分页和无限滚动两种页面模式'),
        ('📥 结果导出', '一键导出符合条件的候选人 CSV 文件'),
        ('🔄 自动推荐', '可将合适人选自动推荐给用人团队')
    ]
    
    for icon_text, desc in features:
        p = add_paragraph_custom(doc, f'{icon_text}：{desc}', size=11)
        p.paragraph_format.left_indent = Inches(0.3)
    
    doc.add_page_break()
    
    # 二、安装步骤
    add_heading_custom(doc, '二、安装步骤', 1)
    
    add_heading_custom(doc, '步骤 1：准备插件文件', 2)
    add_paragraph_custom(doc, '确保你拥有以下文件：')
    
    files = [
        'manifest.json', 'content.js', 'sidebar.html', 'sidebar.js',
        'background.js', 'config.js', 'html2canvas.min.js',
        'images/ 文件夹（包含图标文件）'
    ]
    for f in files:
        p = add_paragraph_custom(doc, f'• {f}')
        p.paragraph_format.left_indent = Inches(0.3)
    
    add_heading_custom(doc, '步骤 2：在 Chrome 中加载插件', 2)
    
    steps = [
        '打开 Chrome 浏览器，在地址栏输入：chrome://extensions/',
        '打开右上角的"开发者模式"开关',
        '点击左上角的"加载已解压的扩展程序"按钮',
        '选择包含插件文件的文件夹',
        '插件安装成功后，会在扩展程序列表中显示"Moka简历筛选Agent"'
    ]
    for i, step in enumerate(steps, 1):
        add_paragraph_custom(doc, f'{i}. {step}')
    
    add_heading_custom(doc, '步骤 3：固定插件图标', 2)
    
    steps2 = [
        '点击 Chrome 工具栏的"扩展程序"图标（拼图形状）',
        '找到"Moka简历筛选Agent"',
        '点击"固定"按钮，将插件图标固定到工具栏'
    ]
    for i, step in enumerate(steps2, 1):
        add_paragraph_custom(doc, f'{i}. {step}')
    
    doc.add_page_break()
    
    # 三、界面说明
    add_heading_custom(doc, '三、界面说明', 1)
    
    add_paragraph_custom(doc, '点击浏览器工具栏的插件图标，会打开侧边栏面板，界面如下图所示：')
    
    # 界面结构表格
    add_heading_custom(doc, '界面区域说明', 2)
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = ['区域', '功能说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    rows_data = [
        ('状态栏', '显示当前筛选状态和进度（已筛选/目标数量）'),
        ('大模型配置', '配置 AI 模型的 API 信息'),
        ('筛选配置', '设置筛选标准和目标数量'),
        ('操作控制', '开始、暂停、停止筛选，导出结果'),
        ('实时日志', '显示筛选过程的详细日志信息')
    ]
    
    for i, (area, desc) in enumerate(rows_data, 1):
        table.rows[i].cells[0].text = area
        table.rows[i].cells[1].text = desc
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # 四、配置说明
    add_heading_custom(doc, '四、配置说明', 1)
    
    add_heading_custom(doc, '4.1 大模型配置', 2)
    add_paragraph_custom(doc, '首次使用需要配置大模型 API 信息：')
    
    # 配置表格
    table2 = doc.add_table(rows=4, cols=3)
    table2.style = 'Light Grid Accent 1'
    
    headers2 = ['配置项', '说明', '示例']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    config_data = [
        ('模型类型', '使用的大模型名称', 'doubao-seed-1-6-lite'),
        ('API Key', '你的 API 密钥', '从模型服务商获取'),
        ('Base URL', 'API 请求地址', 'https://ark.cn-beijing.volces.com/api/v3')
    ]
    
    for i, (item, desc, example) in enumerate(config_data, 1):
        table2.rows[i].cells[0].text = item
        table2.rows[i].cells[1].text = desc
        table2.rows[i].cells[2].text = example
        for cell in table2.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    add_paragraph_custom(doc, '')
    add_paragraph_custom(doc, '支持的模型服务商：', bold=True)
    services = ['火山引擎（豆包大模型）', 'OpenAI', '其他兼容 OpenAI API 格式的服务商']
    for s in services:
        p = add_paragraph_custom(doc, f'• {s}')
        p.paragraph_format.left_indent = Inches(0.3)
    
    add_heading_custom(doc, '4.2 筛选配置', 2)
    
    add_paragraph_custom(doc, '初筛标准：', bold=True)
    add_paragraph_custom(doc, '用于第一轮快速筛选，判断候选人是否值得进一步查看简历。')
    examples = ['985院校，理工专业', '本科及以上学历，计算机相关专业', '3年以上工作经验']
    for ex in examples:
        p = add_paragraph_custom(doc, f'• {ex}')
        p.paragraph_format.left_indent = Inches(0.3)
    
    add_paragraph_custom(doc, '')
    add_paragraph_custom(doc, '完整筛选标准：', bold=True)
    add_paragraph_custom(doc, '用于详细评估候选人简历，做出最终筛选决定。')
    examples2 = [
        '211或者QS前100院校，并且有算法相关经验',
        '有互联网大厂工作经验，熟悉Java/Python',
        '硕士学历，发表过顶会论文'
    ]
    for ex in examples2:
        p = add_paragraph_custom(doc, f'• {ex}')
        p.paragraph_format.left_indent = Inches(0.3)
    
    doc.add_page_break()
    
    # 五、使用流程
    add_heading_custom(doc, '五、使用流程', 1)
    
    add_heading_custom(doc, '5.1 准备工作', 2)
    prep_steps = [
        '登录 Moka 系统：在浏览器中打开 Moka 招聘系统并登录',
        '进入人才库：导航到需要筛选的候选人列表页面',
        '打开插件：点击 Chrome 工具栏的插件图标'
    ]
    for i, step in enumerate(prep_steps, 1):
        add_paragraph_custom(doc, f'{i}. {step}')
    
    add_heading_custom(doc, '5.2 开始筛选', 2)
    add_paragraph_custom(doc, '1. 检查配置：确保大模型配置和筛选配置已正确填写')
    add_paragraph_custom(doc, '2. 点击开始：点击"开始筛选"按钮')
    add_paragraph_custom(doc, '3. 等待处理：插件会自动执行以下操作：')
    
    auto_steps = [
        '提取当前页面候选人信息',
        '使用 AI 进行初筛',
        '对通过初筛的候选人进入详情页',
        '抓取简历图片并进行详细评估',
        '自动翻页继续处理'
    ]
    for step in auto_steps:
        p = add_paragraph_custom(doc, f'   • {step}')
        p.paragraph_format.left_indent = Inches(0.5)
    
    add_heading_custom(doc, '5.3 监控进度', 2)
    add_paragraph_custom(doc, '• 在"状态栏"查看当前进度')
    add_paragraph_custom(doc, '• 在"实时日志"区域查看详细处理信息')
    add_paragraph_custom(doc, '• 可以随时点击"暂停"按钮暂停筛选')
    
    add_heading_custom(doc, '5.4 导出结果', 2)
    add_paragraph_custom(doc, '筛选完成后，点击"导出结果 CSV"按钮，系统会自动下载以下文件：')
    
    p = add_paragraph_custom(doc, '• 合格候选人_YYYY-MM-DD.csv - 符合筛选条件的候选人')
    p.paragraph_format.left_indent = Inches(0.3)
    p = add_paragraph_custom(doc, '• API失败候选人_YYYY-MM-DD.csv - API 调用失败的候选人')
    p.paragraph_format.left_indent = Inches(0.3)
    
    doc.add_page_break()
    
    # 六、操作按钮说明
    add_heading_custom(doc, '六、操作按钮说明', 1)
    
    table3 = doc.add_table(rows=6, cols=3)
    table3.style = 'Light Grid Accent 1'
    
    headers3 = ['按钮', '功能', '使用场景']
    for i, header in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    button_data = [
        ('开始筛选', '启动筛选流程', '配置完成后开始筛选'),
        ('暂停', '暂停当前筛选', '需要临时中断时'),
        ('继续', '恢复筛选', '暂停后恢复'),
        ('停止', '完全停止筛选', '需要终止筛选时'),
        ('导出结果 CSV', '下载筛选结果', '筛选完成或需要中途导出')
    ]
    
    for i, (btn, func, scene) in enumerate(button_data, 1):
        table3.rows[i].cells[0].text = btn
        table3.rows[i].cells[1].text = func
        table3.rows[i].cells[2].text = scene
        for cell in table3.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_page_break()
    
    # 七、常见问题
    add_heading_custom(doc, '七、常见问题', 1)
    
    questions = [
        ('Q1: 插件无法启动怎么办？', [
            '检查是否已正确安装插件（开发者模式是否开启）',
            '刷新 Moka 页面后重试',
            '检查 Chrome 控制台是否有错误信息'
        ]),
        ('Q2: API 调用失败怎么办？', [
            'API Key 错误：检查 API Key 是否正确',
            '网络问题：检查网络连接',
            '模型名称错误：确认模型类型名称正确',
            '余额不足：检查 API 账户余额'
        ]),
        ('Q3: 无法获取简历图片？', [
            '确保候选人详情页已完全加载',
            '检查简历是否正常显示',
            '尝试刷新页面后重新开始'
        ]),
        ('Q4: 筛选结果不准确？', [
            '优化筛选标准：使用更具体、明确的描述',
            '调整模型参数：尝试不同的模型',
            '分批筛选：减少单次筛选数量，提高准确性'
        ]),
        ('Q5: 自动推荐功能无效？', [
            '确认已勾选"直接推荐合适人选给用人团队"',
            '确认已填写负责人姓名',
            '确认当前页面支持推荐功能（无限滚动页面）',
            '检查候选人是否已被其他岗位锁定'
        ])
    ]
    
    for q, answers in questions:
        add_paragraph_custom(doc, q, bold=True, size=12)
        for ans in answers:
            p = add_paragraph_custom(doc, f'• {ans}')
            p.paragraph_format.left_indent = Inches(0.3)
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # 八、注意事项
    add_heading_custom(doc, '八、注意事项', 1)
    
    notices = [
        '页面保持：筛选过程中请勿关闭或刷新 Moka 页面',
        '网络稳定：确保网络连接稳定，避免 API 调用失败',
        'API 限额：注意 API 调用次数限制和费用',
        '隐私保护：妥善保管 API Key，不要分享给他人',
        '结果复核：AI 筛选结果仅供参考，建议人工复核重要岗位'
    ]
    
    for notice in notices:
        p = add_paragraph_custom(doc, f'⚠️ {notice}')
    
    # 九、技术支持
    add_heading_custom(doc, '九、技术支持', 1)
    add_paragraph_custom(doc, '如有问题或建议，请联系开发团队。')
    
    # 页脚
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('版本：2.0    更新日期：2026-02-25')
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 保存文档
    doc.save('人选筛选助手使用说明.docx')
    print('Word 文档已生成：人选筛选助手使用说明.docx')

if __name__ == '__main__':
    create_doc()
